from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement


def set_run_font_size(run, size):
    run_font = run._element.find('.//w:rPr/w:rFonts')
    sz_element = OxmlElement('w:sz')
    sz_element.set('w:val', str(size))
    run_font.append(sz_element)


def markdown_to_docx(text):
    lines = text.split('\n')
    doc = Document()
    doc.add_heading(lines[0], level=1)

    for line in lines[1:]:
        if line.startswith('# '):
            doc.add_heading(line[2:], level=2)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=3)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=4)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=5)
        elif line.startswith('##### '):
            doc.add_heading(line[6:], level=6)
        elif line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. '):
            doc.add_paragraph(line[3:], style='List Number')
        else:
            p = doc.add_paragraph()
            for word in line.split():
                if word.startswith('**'):
                    run = p.add_run(word[2:] + ' ')
                    run.bold = True
                elif word.startswith('*'):
                    run = p.add_run(word[1:] + ' ')
                    run.italic = True
                elif word.startswith('__'):
                    run = p.add_run(word[2:] + ' ')
                    run.bold = True
                elif word.startswith('_'):
                    run = p.add_run(word[1:] + ' ')
                    run.italic = True
                else:
                    run = p.add_run(word + ' ')
                set_run_font_size(run, Pt(12))

    doc.save('res.docx')

