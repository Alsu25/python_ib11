import sys
from docx import Document
d = Document()
a = input().lower().strip()
b = input().lower().strip()
for i in sys.stdin:
    i = i.strip()
    if i.replace(' ', '').isalpha() == True:
        d.add_heading('Приглашение на мероприятие', 0)
        p = d.add_paragraph('')
        p.add_run(f'{i}, приглашаем вас на мероприятие, которое состоится {a} {b}. \n')
        p.add_run('Спасибо за внимание!')
        d.add_page_break()
d.save('test.docx')